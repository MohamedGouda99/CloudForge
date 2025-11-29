from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Create document
doc = Document()

# Set up styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('CloudForge Project Handover Document', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Metadata
doc.add_paragraph(f'Handover Date: {datetime.now().strftime("%B %d, %Y")}')
doc.add_paragraph('Project: CloudForge - Diagram-First Terraform Authoring Platform')
doc.add_paragraph('Repository: https://github.com/MohamedGouda99/CloudForge.git')

doc.add_page_break()

# Table of Contents
doc.add_heading('Table of Contents', 1)
toc_items = [
    '1. Project Overview',
    '2. System Architecture',
    '3. Technology Stack',
    '4. Environment Setup',
    '5. Project Structure',
    '6. Key Components Deep Dive',
    '7. Recent Development Work',
    '8. Configuration & Environment Variables',
    '9. Running the Application',
    '10. Common Operations',
    '11. Troubleshooting Guide',
    '12. Development Workflow',
    '13. Known Issues & Future Work',
    '14. Contact & Resources'
]
for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# 1. Project Overview
doc.add_heading('1. Project Overview', 1)
doc.add_paragraph(
    'CloudForge is a diagram-first Infrastructure as Code (IaC) platform that allows users to '
    'visually design cloud infrastructure using a drag-and-drop interface and automatically generate '
    'Terraform configuration files. The platform supports AWS, Azure, and GCP resources.'
)

doc.add_heading('Key Features:', 2)
features = [
    'Visual Infrastructure Designer: Drag-and-drop canvas for creating cloud architecture diagrams',
    'Multi-Cloud Support: AWS, Azure, and GCP resources with official cloud provider icons',
    'Terraform Generation: Automatic conversion of diagrams to production-ready Terraform code',
    'Resource Configuration: Detailed configuration modals for each cloud resource type',
    'Dark Mode Support: Full UI support for light and dark themes',
    'Real-time Validation: Terraform validation before deployment',
    'Project Management: Save, load, and version control infrastructure designs'
]
for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('Use Cases:', 2)
use_cases = [
    'Rapid prototyping of cloud infrastructure',
    'Infrastructure documentation and visualization',
    'Learning Terraform and cloud architecture',
    'Collaborative infrastructure design',
    'Infrastructure as Code generation without writing HCL manually'
]
for uc in use_cases:
    doc.add_paragraph(uc, style='List Bullet')

doc.add_page_break()

# 2. System Architecture
doc.add_heading('2. System Architecture', 1)
doc.add_paragraph(
    'CloudForge follows a modern microservices architecture with clear separation of concerns:'
)

# Architecture table
table = doc.add_table(rows=7, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Component'
hdr_cells[1].text = 'Technology'
hdr_cells[2].text = 'Port'
hdr_cells[3].text = 'Purpose'

data = [
    ('Frontend', 'React 18 + Vite', '3000', 'Visual diagram editor UI'),
    ('Backend API', 'FastAPI + Python 3.11', '8000', 'REST API, Terraform generation'),
    ('PostgreSQL', 'PostgreSQL 15', '5432', 'Project & user data storage'),
    ('Redis', 'Redis 7', '6379', 'Caching & session management'),
    ('Celery Worker', 'Celery + Redis', 'N/A', 'Background task processing'),
    ('Nginx', 'Nginx (prod only)', '80/443', 'Reverse proxy & static serving')
]

for i, (comp, tech, port, purpose) in enumerate(data, 1):
    row = table.rows[i].cells
    row[0].text = comp
    row[1].text = tech
    row[2].text = port
    row[3].text = purpose

doc.add_paragraph()
doc.add_heading('Data Flow:', 2)
flow_steps = [
    'User creates diagram in React frontend using ReactFlow library',
    'Diagram state (nodes, edges, configurations) stored in component state',
    'User triggers "Generate Terraform" action',
    'Frontend sends diagram JSON to FastAPI backend via REST API',
    'Backend parses diagram structure and resource configurations',
    'Terraform generator creates .tf files based on resource types and relationships',
    'Generated files returned to frontend or stored in database',
    'User can download, preview, or deploy generated Terraform code'
]
for i, step in enumerate(flow_steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

doc.add_page_break()

# Save document
doc.save('CloudForge_Handover_Document.docx')
print("OK - Document created successfully!")
